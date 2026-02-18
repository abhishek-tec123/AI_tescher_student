import os
import json
import logging
import requests
import fitz
import pytesseract
import docx
import pandas as pd

from io import BytesIO, StringIO, IOBase
from urllib.parse import urlparse
from typing import List, Union
from concurrent.futures import ThreadPoolExecutor, ProcessPoolExecutor, as_completed
from bs4 import BeautifulSoup
from PIL import Image
from langchain_core.documents import Document
from langchain_core.document_loaders.base import BaseLoader

# =========================================================
# Logging Setup
# =========================================================
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s | %(levelname)s | %(message)s",
)
logger = logging.getLogger(__name__)


# =========================================================
# Parallel OCR PDF Extraction
# =========================================================
from embaddings.imgae_basePdf import extract_img_pdf_text_parallel
# =========================================================
# Mixed File Type Loader
# =========================================================
class MixedFileTypeLoader(BaseLoader):
    """
    Load documents from file paths or URLs.
    Supports: PDF (OCR enabled), DOCX, DOC, TXT, MD, HTML, CSV, JSON
    """

    SUPPORTED_EXTENSIONS = {
        ".pdf", ".docx", ".doc", ".txt",
        ".md", ".html", ".htm", ".csv", ".json"
    }

    def __init__(self, file_paths: Union[str, List[str]], verbose: bool = True):
        if isinstance(file_paths, str):
            file_paths = [file_paths]

        self.file_paths = self._filter_supported_files(file_paths)
        self.verbose = verbose

        if not verbose:
            logger.setLevel(logging.WARNING)

        if not self.file_paths:
            raise ValueError("No valid files provided.")

    # -----------------------------------------------------
    # Helpers
    # -----------------------------------------------------
    def _is_url(self, path: str) -> bool:
        return path.startswith("http://") or path.startswith("https://")

    def _filter_supported_files(self, paths: List[str]) -> List[str]:
        supported = []

        for path in paths:
            if self._is_url(path):
                ext = os.path.splitext(urlparse(path).path)[1].lower()
                if ext in self.SUPPORTED_EXTENSIONS:
                    supported.append(path)
                else:
                    logger.warning(f"Unsupported URL file type: {path}")

            elif os.path.isfile(path):
                ext = os.path.splitext(path)[1].lower()
                if ext in self.SUPPORTED_EXTENSIONS:
                    supported.append(path)
                else:
                    logger.warning(f"Unsupported file type: {path}")
            else:
                logger.warning(f"Invalid path: {path}")

        return supported

    # -----------------------------------------------------
    # Main Load Method
    # -----------------------------------------------------
    def load(self) -> List[Document]:
        documents = []

        with ThreadPoolExecutor() as executor:
            futures = [
                executor.submit(self._load_single, path)
                for path in self.file_paths
            ]

            for future in as_completed(futures):
                result = future.result()
                if result:
                    if isinstance(result, list):
                        documents.extend(result)
                    else:
                        documents.append(result)

        return documents

    # -----------------------------------------------------
    # Single File Loader
    # -----------------------------------------------------
    def _load_single(self, path: str) -> Union[Document, List[Document], None]:
        try:
            is_url = self._is_url(path)
            ext = os.path.splitext(
                urlparse(path).path if is_url else path
            )[1].lower()

            # PDF special handling (OCR-enabled)
            if ext == ".pdf" and not is_url:
                text = extract_img_pdf_text_parallel(path)
                return Document(page_content=text, metadata={"source": path})

            loader_fn = {
                ".docx": self._load_word,
                ".doc": self._load_word,
                ".txt": self._load_txt,
                ".md": self._load_txt,
                ".html": self._load_html,
                ".htm": self._load_html,
                ".csv": self._load_csv,
                ".json": self._load_json,
                ".pdf": self._load_pdf_url  # PDF from URL only
            }.get(ext)

            if not loader_fn:
                return None

            file_buffer = self._get_file_buffer(path)
            result = loader_fn(file_buffer)

            if isinstance(result, list):
                for doc in result:
                    doc.metadata["source"] = path
                return result

            return Document(page_content=result, metadata={"source": path})

        except Exception as e:
            logger.error(f"Error processing {path}: {e}", exc_info=True)
            return None

    # -----------------------------------------------------
    # File Buffer
    # -----------------------------------------------------
    def _get_file_buffer(self, path: str) -> Union[IOBase, StringIO]:
        if self._is_url(path):
            response = requests.get(path, timeout=20)
            response.raise_for_status()
            return BytesIO(response.content)

        ext = os.path.splitext(path)[1].lower()

        if ext in {".pdf", ".docx", ".doc"}:
            return open(path, "rb")

        try:
            return open(path, "r", encoding="utf-8")
        except UnicodeDecodeError:
            return open(path, "r", encoding="latin-1")

    # -----------------------------------------------------
    # File Type Loaders
    # -----------------------------------------------------
    def _load_pdf_url(self, file: BytesIO) -> str:
        file.seek(0)
        with fitz.open(stream=file.read(), filetype="pdf") as doc:
            return "\n".join(page.get_text() for page in doc)

    def _load_word(self, file: BytesIO) -> str:
        file.seek(0)
        document = docx.Document(file)
        return "\n".join(p.text for p in document.paragraphs if p.text.strip())

    def _load_txt(self, file: StringIO) -> str:
        file.seek(0)
        return file.read()

    def _load_html(self, file: StringIO) -> str:
        file.seek(0)
        soup = BeautifulSoup(file.read(), "html.parser")
        return soup.get_text(separator="\n", strip=True)

    def _load_csv(self, file: Union[StringIO, IOBase]) -> List[Document]:
        file.seek(0)
        df = pd.read_csv(file)
        return [
            Document(
                page_content="\n".join(f"{col}: {row[col]}" for col in df.columns),
                metadata={"row": idx}
            )
            for idx, row in df.iterrows()
        ]

    def _load_json(self, file: Union[StringIO, IOBase]) -> List[Document]:
        file.seek(0)
        data = json.load(file)

        if isinstance(data, list):
            return [
                Document(page_content=json.dumps(item), metadata={})
                for item in data
            ]

        return [Document(page_content=json.dumps(data), metadata={})]


# # =========================================================
# # Example Usage
# # =========================================================
# if __name__ == "__main__":

#     files = [
#         # "/Users/macbook/Desktop/langchain_adk/data/BPSC Computer TRE 4.0_____ Part-1 (4 sheet) final.pdf",
#         "/Users/macbook/Desktop/langchain_adk/data/ML+Cheat+Sheet_2.pdf"
#     ]

#     loader = MixedFileTypeLoader(files)
#     documents = loader.load()

#     print(f"\nLoaded {len(documents)} documents\n")
#     print(documents[0].page_content)


# from utility import load_documents, split_documents, embed_chunks,build_embedding_json_for_db

# def main():
#     # 📂 Step 1: Define input source(s)
#     file_inputs = ["./docs"]  # Replace with your file paths or folder

#     # ⚙️ Step 2: Load documents
#     print("[*] Loading documents...")
#     docs = load_documents(file_inputs)

#     # 📏 Step 3: Split documents into manageable chunks
#     print(f"[*] Splitting {len(docs)} documents into chunks...")
#     chunks = split_documents(docs, chunk_size=1000, chunk_overlap=200)

#     # 🤖 Step 4: Load embedding model
#     print("[*] Loading embedding model...")
#     model_name = "sentence-transformers/all-MiniLM-L12-v2"
#     embedding_model = HuggingFaceEmbeddings(
#         model_name=model_name,
#         model_kwargs={"device": "cpu"},  # Use "cuda" for GPU
#         encode_kwargs={"normalize_embeddings": True}
#     )

#     # 🔢 Step 5: Embed the chunks
#     print(f"[*] Embedding {len(chunks)} chunks...")
#     embeddings = embed_chunks(chunks, embedding_model)

#     # 🧱 Step 6: Build embedding data for DB
#     print("[*] Building embedding JSON for DB...")
#     embedding_json, doc_ids = build_embedding_json_for_db(
#         chunks, embeddings, embedding_model_name=model_name
#     )

#     print(f"[✅] Processed {len(embedding_json)} embeddings from {len(doc_ids)} documents.")
#     return embedding_json  # <-- Return instead of saving

# # Example usage
# if __name__ == "__main__":
#     output = main()
#     import json
#     print(json.dumps(output, indent=2, ensure_ascii=False))